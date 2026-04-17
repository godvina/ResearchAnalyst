"""Convert the v2 success plan to Word."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT = "docs/investigative-intelligence-success-plan-v2.md"
OUTPUT = "docs/Investigative-Intelligence-PoC-Success-Plan-v2.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# Title page
for _ in range(5):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("[Customer Name]")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Investigative Intelligence Platform")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
run.bold = True
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("PoC Success Plan")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x48, 0xBB, 0x78)
doc.add_paragraph()
for line in ["Emerging Tech Solutions", "Amazon Web Services", "April 2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
doc.add_page_break()

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

def add_table_from_lines(lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    header = rows[0]
    data_rows = [r for r in rows[1:] if not all(set(c) <= {"-", " ", ":"} for c in r)]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell_text in enumerate(header):
        table.rows[0].cells[i].text = cell_text
    for r_idx, row in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < len(table.columns):
                table.rows[r_idx + 1].cells[c_idx].text = cell_text

with open(INPUT, "r", encoding="utf-8") as f:
    content = f.read()

lines = content.split("\n")
start_idx = 0
for idx, line in enumerate(lines):
    if line.startswith("## PoC Team"):
        start_idx = idx
        break

i = start_idx
table_buffer = []
in_table = False
in_code = False

while i < len(lines):
    line = lines[i]
    if line.strip().startswith("```"):
        if in_code:
            in_code = False
            i += 1
            continue
        else:
            in_code = True
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            in_code = False
            i += 1
            continue
    if "|" in line and not in_code:
        if not in_table:
            in_table = True
            table_buffer = []
        table_buffer.append(line)
        i += 1
        continue
    elif in_table:
        add_table_from_lines(table_buffer)
        doc.add_paragraph()
        table_buffer = []
        in_table = False
    if line.startswith("## "):
        add_heading(line[3:].strip(), level=1)
    elif line.startswith("### "):
        add_heading(line[4:].strip(), level=2)
    elif line.startswith("#### "):
        add_heading(line[5:].strip(), level=3)
    elif line.startswith("---"):
        pass
    elif line.strip() == "":
        pass
    elif re.match(r"^\s*[-•] ", line):
        text = re.sub(r"^\s*[-•] ", "", line).strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        indent = len(line) - len(line.lstrip())
        style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
        doc.add_paragraph(text, style=style_name)
    elif re.match(r"^\d+\.\s", line):
        text = re.sub(r"^\d+\.\s*", "", line).strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        doc.add_paragraph(text, style="List Number")
    elif line.startswith("**") and line.endswith("**"):
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*").strip())
        run.bold = True
    elif line.startswith("*") and line.endswith("*"):
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*").strip())
        run.italic = True
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    else:
        clean = line.strip()
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)
        clean = re.sub(r"`(.+?)`", r"\1", clean)
        if clean:
            doc.add_paragraph(clean)
    i += 1

if in_table and table_buffer:
    add_table_from_lines(table_buffer)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
