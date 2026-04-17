"""Generate Investigative Intelligence PoC Word documents.

Produces two .docx files:
1. CIO Planning Call Questionnaire
2. PoC Success Plan Template

Both are enriched with real architecture knowledge from the Research Analyst
platform we built (S3, Neptune, OpenSearch, Aurora, Bedrock, Step Functions).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def style_header_row(table, bg="1B2A4A"):
    """Style the first row of a table as a dark header."""
    for cell in table.rows[0].cells:
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)


def add_table_row(table, cells, shade=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = str(text)
        for p in row.cells[i].paragraphs:
            p.style.font.size = Pt(9)
        if shade:
            set_cell_shading(row.cells[i], shade)
    return row


def make_questionnaire():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # -- Title --
    title = doc.add_heading("Investigative Intelligence Platform", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("CIO Planning Call — Scoping Questionnaire", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Purpose: Scope a 500TB pilot for an AI-powered investigative intelligence "
        "platform. This questionnaire captures the information needed to produce a "
        "2-page Project Statement (Objective, Scope, Approach, Timeline, Price)."
    )
    doc.add_paragraph("Meeting Duration: 1 hour  |  Format: Take notes in the Answer column")
    doc.add_paragraph("")

    # ---- Section 1: Business Objectives ----
    doc.add_heading("1. Business Objectives (10 min)", level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.columns[0].width = Cm(1)
    t.columns[1].width = Cm(9)
    t.columns[2].width = Cm(7)
    t.rows[0].cells[0].text = "#"
    t.rows[0].cells[1].text = "Question"
    t.rows[0].cells[2].text = "Answer"
    style_header_row(t)

    biz_qs = [
        ("1.1", "What is the primary mission outcome for this pilot? (faster case resolution, cross-case pattern detection, analyst productivity, leadership visibility)"),
        ("1.2", "Who are the end users? (investigators, analysts, prosecutors, leadership) Approximate number?"),
        ("1.3", "What does success look like? What would make you say 'let's go to production'?"),
        ("1.4", "Is there a specific case type or investigation you want to pilot first?"),
        ("1.5", "Top 3 capabilities from the demo that matter most? (cross-case analysis, AI summaries/briefings, pattern discovery, entity graph, geospatial mapping, semantic search)"),
        ("1.6", "Is there a timeline or event driving urgency? (budget cycle, leadership review, active investigation)"),
    ]
    for num, q in biz_qs:
        add_table_row(t, [num, q, ""])

    # ---- Section 2: Data Landscape ----
    doc.add_heading("2. Data Landscape (15 min)", level=2)
    doc.add_paragraph(
        "These questions directly impact architecture sizing, ingestion pipeline design, "
        "and whether we need OCR (Textract) vs. text extraction (PyPDF2). From our build "
        "experience: scanned PDFs without a text layer require Textract at ~$1.50/1000 pages, "
        "while readable PDFs process at ~$0.003/doc via Bedrock entity extraction."
    ).italic = True

    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.columns[0].width = Cm(1)
    t2.columns[1].width = Cm(9)
    t2.columns[2].width = Cm(7)
    t2.rows[0].cells[0].text = "#"
    t2.rows[0].cells[1].text = "Question"
    t2.rows[0].cells[2].text = "Answer"
    style_header_row(t2)

    data_qs = [
        ("2.1", "Of the 500TB, what file types? (PDF, Word, Excel, images, email .msg/.eml, structured DB exports, other)"),
        ("2.2", "Are PDFs machine-readable (text layer) or scanned images requiring OCR? Rough % split?"),
        ("2.3", "Is data already in S3 or needs migration? From where? (on-prem, other cloud, network shares)"),
        ("2.4", "Is there an existing hierarchy? (Matter → Case → Collection → Document, or flat files?)"),
        ("2.5", "Existing metadata? (case ID, date, custodian, classification level, sensitivity labels)"),
        ("2.6", "Data sensitivity level? (CUI, classified, law enforcement sensitive, PII content)"),
        ("2.7", "Existing entity data? (people, orgs, locations already tagged in a database or case management system)"),
        ("2.8", "How frequently does new data arrive? (batch daily, streaming, ad-hoc uploads)"),
        ("2.9", "For the pilot, can we start with a representative subset? What size? (1TB, 10TB, 50TB)"),
    ]
    for num, q in data_qs:
        add_table_row(t2, [num, q, ""])

    # ---- Section 3: Technical Environment ----
    doc.add_heading("3. Technical Environment (10 min)", level=2)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    t3.columns[0].width = Cm(1)
    t3.columns[1].width = Cm(9)
    t3.columns[2].width = Cm(7)
    t3.rows[0].cells[0].text = "#"
    t3.rows[0].cells[1].text = "Question"
    t3.rows[0].cells[2].text = "Answer"
    style_header_row(t3)

    tech_qs = [
        ("3.1", "GovCloud region? Existing account or new? (us-gov-west-1 / us-gov-east-1)"),
        ("3.2", "AWS services already in use? (S3, VPC, IAM, existing databases, existing OpenSearch)"),
        ("3.3", "Existing VPC and networking we need to integrate with? VPN/Direct Connect to on-prem?"),
        ("3.4", "Authentication requirements? (PIV/CAC, SAML/OIDC, Active Directory)"),
        ("3.5", "FedRAMP, FISMA, or ATO requirements for the pilot environment?"),
        ("3.6", "Data residency or encryption requirements beyond AWS defaults? (CMK, FIPS 140-2)"),
    ]
    for num, q in tech_qs:
        add_table_row(t3, [num, q, ""])

    # ---- Section 4: Team & Implementation ----
    doc.add_heading("4. Team & Implementation Model (10 min)", level=2)
    t4 = doc.add_table(rows=1, cols=3)
    t4.style = "Table Grid"
    t4.columns[0].width = Cm(1)
    t4.columns[1].width = Cm(9)
    t4.columns[2].width = Cm(7)
    t4.rows[0].cells[0].text = "#"
    t4.rows[0].cells[1].text = "Question"
    t4.rows[0].cells[2].text = "Answer"
    style_header_row(t4)

    team_qs = [
        ("4.1", "Do you have a technical team to implement, or need ProServe / partner?"),
        ("4.2", "If internal — how many engineers? AWS experience? Python? Graph databases?"),
        ("4.3", "Preferred model? (AWS builds, your team builds with guidance, hybrid)"),
        ("4.4", "Who is the technical decision maker and day-to-day POC?"),
        ("4.5", "Existing tools to integrate with? (case management, BI, SIEM, eDiscovery)"),
    ]
    for num, q in team_qs:
        add_table_row(t4, [num, q, ""])

    # ---- Section 5: Budget & Timeline ----
    doc.add_heading("5. Budget & Timeline (10 min)", level=2)
    t5 = doc.add_table(rows=1, cols=3)
    t5.style = "Table Grid"
    t5.columns[0].width = Cm(1)
    t5.columns[1].width = Cm(9)
    t5.columns[2].width = Cm(7)
    t5.rows[0].cells[0].text = "#"
    t5.rows[0].cells[1].text = "Question"
    t5.rows[0].cells[2].text = "Answer"
    style_header_row(t5)

    budget_qs = [
        ("5.1", "~$200K budget — combined ProServe + AWS consumption, or separate?"),
        ("5.2", "Flexibility if pilot scope requires more? Or hard cap?"),
        ("5.3", "Desired pilot start date and duration? (8 weeks, 12 weeks)"),
        ("5.4", "When do you need a decision/readout to leadership?"),
        ("5.5", "Path to production funding if pilot succeeds? Rough timeline?"),
    ]
    for num, q in budget_qs:
        add_table_row(t5, [num, q, ""])

    # ---- Section 6: Data Governance & Access Control ----
    doc.add_heading("6. Data Governance & Access Control (10 min)", level=2)
    doc.add_paragraph(
        "These questions determine the access control model (label-based vs. row-level security), "
        "authentication integration, and compliance requirements. The platform supports both "
        "document-level classification labels and PostgreSQL row-level security for team isolation."
    ).italic = True

    t6g = doc.add_table(rows=1, cols=3)
    t6g.style = "Table Grid"
    t6g.columns[0].width = Cm(1)
    t6g.columns[1].width = Cm(9)
    t6g.columns[2].width = Cm(7)
    t6g.rows[0].cells[0].text = "#"
    t6g.rows[0].cells[1].text = "Question"
    t6g.rows[0].cells[2].text = "Answer"
    style_header_row(t6g)

    gov_qs = [
        ("6.1", "How many distinct teams or task forces will use the platform simultaneously? (single team vs. multiple teams with case isolation)"),
        ("6.2", "Do teams need to be prevented from seeing each other's cases, or is cross-case visibility acceptable?"),
        ("6.3", "What roles exist in your investigative workflow? (Lead Investigator, Analyst, Prosecutor, Read-Only Reviewer, Administrator, other)"),
        ("6.4", "Does your agency use a formal classification hierarchy for evidence? (UNCLASSIFIED, CUI, SECRET, TOP SECRET, or custom)"),
        ("6.5", "Can classification labels change over time? (e.g., evidence reclassified during investigation)"),
        ("6.6", "Are there documents that specific individuals must never see, regardless of clearance level? (conflict-of-interest exclusions)"),
        ("6.7", "What identity provider does your agency use? (AWS SSO, Okta, Azure AD, PIV/CAC via SAML, other)"),
        ("6.8", "Do users authenticate with PIV/CAC smart cards?"),
        ("6.9", "Is there an existing user directory to integrate with, or should the platform manage its own user registry?"),
        ("6.10", "What audit retention requirements apply? (e.g., 7 years for federal records)"),
        ("6.11", "Do you require audit logs to be tamper-evident? (signed, write-once storage)"),
        ("6.12", "Are there specific NIST 800-53 controls that must be documented?"),
        ("6.13", "What is the expected data retention period for ingested evidence?"),
        ("6.14", "Is there a requirement to purge or redact data after a case is closed?"),
        ("6.15", "Can AI-generated analysis (theories, case files, entity extractions) be deleted independently of source evidence?"),
    ]
    for num, q in gov_qs:
        add_table_row(t6g, [num, q, ""])

    # ---- Section 7: Scope Priorities ----
    doc.add_heading("7. Pilot Scope Priorities (5 min)", level=2)
    doc.add_paragraph("Rank 1-5 for the pilot (1 = must have, 5 = nice to have):")
    t6 = doc.add_table(rows=1, cols=3)
    t6.style = "Table Grid"
    t6.rows[0].cells[0].text = "Capability"
    t6.rows[0].cells[1].text = "AWS Services"
    t6.rows[0].cells[2].text = "Priority"
    style_header_row(t6)

    caps = [
        ("Document ingestion pipeline (S3 → parse → extract → embed → store)", "S3, Step Functions, Lambda, Textract"),
        ("Semantic + keyword search across all documents", "OpenSearch Service"),
        ("Entity extraction & knowledge graph (people, orgs, locations, relationships)", "Bedrock (Claude), Neptune"),
        ("AI case briefings & investigative summaries", "Bedrock (Claude), Aurora"),
        ("Cross-case pattern analysis & entity network discovery", "Neptune, Bedrock"),
        ("Relational metadata store (case hierarchy, document records, entity index)", "Aurora PostgreSQL"),
        ("Geospatial evidence mapping", "OpenSearch, Neptune"),
        ("User access control & multi-tenancy", "IAM, Aurora RLS"),
        ("Custom loader aligned with future OpenSearch managed service", "Lambda, OpenSearch"),
    ]
    for cap, svc in caps:
        add_table_row(t6, [cap, svc, ""])

    # ---- Next Steps ----
    doc.add_heading("Next Steps After This Call", level=2)
    for step in [
        "Compile answers into 2-page Project Statement (Objective, Scope, Approach, Timeline, Price)",
        "Generate PoC Success Plan with evaluation criteria",
        "Schedule technical deep-dive with engineering team",
        "Identify pilot dataset and begin data assessment",
        "Confirm ProServe engagement model and SOW",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    path = os.path.join(OUT_DIR, "Investigative-Intelligence-CIO-Questionnaire.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


def make_success_plan():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Investigative Intelligence Platform", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("PoC Success Plan", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[Customer Name]  |  AWS GovCloud  |  ~$200K Pilot", style="Intense Quote")

    # ---- PoC Team ----
    doc.add_heading("PoC Team", level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "PoC Role"
    t.rows[0].cells[1].text = "Name"
    t.rows[0].cells[2].text = "Organization Role"
    style_header_row(t)
    for role, name, org in [
        ("Business Decision Owner", "", "CIO"),
        ("Technical Decision Owner", "", ""),
        ("Technical Lead", "", ""),
        ("PoC Lead (AWS)", "", ""),
        ("OpenSearch Loader SME", "", "AWS Service Team"),
    ]:
        add_table_row(t, [role, name, org])

    # ---- Timeline ----
    doc.add_heading("PoC Timeline", level=2)
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Target Date"
    t2.rows[0].cells[1].text = "Task"
    t2.rows[0].cells[2].text = "Owner"
    t2.rows[0].cells[3].text = "Status"
    style_header_row(t2)
    for date, task, owner, status in [
        ("Week 1", "Success Plan Finalized", "AWS + Customer", "Not Started"),
        ("Week 1-2", "GovCloud Environment Setup", "Customer", "Not Started"),
        ("Week 2", "Data Assessment & Pilot Dataset Selection", "Joint", "Not Started"),
        ("Week 3-4", "Ingestion Pipeline Build & Data Loading", "ProServe", "Not Started"),
        ("Week 5-6", "Core Platform (Search + Graph + AI Briefings)", "ProServe", "Not Started"),
        ("Week 7", "Cross-Case Analysis & Pattern Discovery", "ProServe", "Not Started"),
        ("Week 8", "User Acceptance Testing", "Customer", "Not Started"),
        ("Week 8", "Executive Readout & Go/No-Go", "Joint", "Not Started"),
    ]:
        add_table_row(t2, [date, task, owner, status])

    # ---- Background ----
    doc.add_heading("Current Situation / Background", level=2)
    for item in [
        "[Customer] has approximately 500TB of investigative case file data in S3",
        "Organization requires AI-powered analytical capabilities across cases",
        "Current tools lack cross-case pattern detection and AI-driven insights",
        "CIO interested in: (1) cross-case analysis, (2) AI summaries, (3) pattern discovery",
        "Deployment target: AWS GovCloud with ~$200K combined ProServe + AWS services",
        "OpenSearch managed loader not yet in GovCloud — custom loader aligned with service team",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ---- Business Drivers ----
    doc.add_heading("Business Drivers", level=2)
    for i, d in enumerate([
        "Accelerate investigative analysis — AI extracts entities, generates case briefings, surfaces connections across millions of documents",
        "Cross-case pattern discovery — Neptune knowledge graph links entities (people, organizations, locations) across cases to find connections analysts wouldn't find manually",
        "Centralized data lake — S3 with OpenSearch provides unified semantic + keyword search across all case data",
        "Reduce time-to-insight — from weeks of manual review to minutes of AI-assisted analysis",
    ], 1):
        doc.add_paragraph(f"{i}. {d}")

    # ---- Architecture ----
    doc.add_heading("Architecture Overview", level=2)
    doc.add_paragraph(
        "Based on a proven architecture built and validated with 8,000+ documents, "
        "2,000+ entities, 1,700+ relationships, and 30+ AI-generated investigative leads. "
        "The following architecture scales to 500TB via Step Functions Distributed Map "
        "and SQS fan-out patterns."
    ).italic = True

    ta = doc.add_table(rows=1, cols=3)
    ta.style = "Table Grid"
    ta.rows[0].cells[0].text = "Component"
    ta.rows[0].cells[1].text = "AWS Service"
    ta.rows[0].cells[2].text = "Purpose"
    style_header_row(ta)
    for comp, svc, purpose in [
        ("Data Lake", "S3", "Raw document storage, centralized access, data lake for all case files"),
        ("Document Search", "OpenSearch Service", "Full-text + semantic vector search across all documents"),
        ("Knowledge Graph", "Neptune", "Entity relationships, cross-case connections, network analysis"),
        ("Relational Store", "Aurora PostgreSQL", "Case metadata, entity index, document records, pgvector embeddings"),
        ("AI / ML", "Bedrock (Claude Haiku)", "Entity extraction, case summarization, pattern analysis, briefing generation"),
        ("OCR", "Textract", "Text extraction from scanned/image PDFs (non-readable documents)"),
        ("Ingestion Pipeline", "Step Functions + Lambda", "Document processing: parse → extract entities → embed → graph load → store"),
        ("Batch Loader", "Lambda + S3 Events", "High-throughput ingestion with cursor-based resumption, cost tracking, blank filtering"),
        ("Custom Loader", "Lambda", "OpenSearch ingestion aligned with future managed service for easy migration"),
    ]:
        add_table_row(ta, [comp, svc, purpose])

    # ---- Ingestion Pipeline Detail ----
    doc.add_heading("Ingestion Pipeline (Proven at Scale)", level=2)
    doc.add_paragraph(
        "The ingestion pipeline processes documents through a Step Functions state machine "
        "with the following stages. This pipeline has been validated with 8,000+ documents "
        "and is designed to scale to millions via Distributed Map (10K+ concurrent executions)."
    )
    for stage in [
        "S3 Event / Batch Trigger → Step Functions execution starts",
        "ResolveConfig → loads case-level and system-level pipeline configuration",
        "ParseDocument → extracts text via PyPDF2 (readable) or Textract (scanned)",
        "ExtractEntities → Bedrock Claude extracts people, organizations, locations, dates, financial amounts, relationships",
        "GenerateEmbedding → Bedrock Titan generates vector embeddings for semantic search",
        "GraphLoad → Neptune bulk loader inserts entities and relationships into knowledge graph",
        "StoreArtifact → Aurora stores document record, entity index, relationship index",
        "OpenSearch Index → full-text + vector index for search",
        "UpdateStatus → marks document as processed, updates case statistics",
    ]:
        doc.add_paragraph(stage, style="List Bullet")

    doc.add_paragraph(
        "Key lessons from our build: VPC-attached Lambdas need 300s+ timeout and 512MB+ memory. "
        "Neptune entity types must match Gremlin query filters exactly. Dual-write to both "
        "Neptune (graph) and Aurora (relational) ensures analytics work even if one store is slow. "
        "Blank/corrupt PDF filtering saves ~40% of Bedrock costs."
    ).italic = True

    # ---- Use Cases ----
    doc.add_heading("Pilot Use Cases", level=2)

    doc.add_heading("Use Case 1: AI Intelligence Briefing", level=3)
    doc.add_paragraph(
        "Analyst selects a case → system generates a comprehensive intelligence assessment. "
        "3-level progressive disclosure: (1) Executive summary with key findings and entity statistics, "
        "(2) Click a finding → detail panel with confidence breakdown, knowledge graph visualization, "
        "and AI justification, (3) View supporting documents with semantic search and highlighted excerpts."
    )

    doc.add_heading("Use Case 2: Cross-Case Pattern Analysis", level=3)
    doc.add_paragraph(
        "Neptune knowledge graph identifies entities (people, organizations, locations) that appear "
        "across multiple cases. Surfaces connections analysts didn't know to look for. Graph visualization "
        "shows entity networks with relationship types, degree centrality scoring, and AI-generated "
        "hypotheses about why connections matter."
    )

    doc.add_heading("Use Case 3: Semantic Document Search", level=3)
    doc.add_paragraph(
        "Natural language queries across all ingested documents. Results ranked by semantic relevance "
        "using vector embeddings (Bedrock Titan → Aurora pgvector or OpenSearch kNN). Supports both "
        "within-case and cross-case search with highlighted excerpts and drill-down to source documents."
    )

    # ---- Key Success Criteria ----
    doc.add_heading("Key Success Criteria", level=2)
    tc = doc.add_table(rows=1, cols=4)
    tc.style = "Table Grid"
    tc.rows[0].cells[0].text = "#"
    tc.rows[0].cells[1].text = "Criteria"
    tc.rows[0].cells[2].text = "Target"
    tc.rows[0].cells[3].text = "PoC Result"
    style_header_row(tc)
    for num, crit, target in [
        ("1", "Ingestion throughput", "1,000+ docs/hour sustained"),
        ("2", "AI briefing quality (analyst rating 1-5)", "≥ 4.0"),
        ("3", "Cross-case pattern accuracy (SME validated)", "≥ 85%"),
        ("4", "Search relevance (precision@10)", "≥ 80%"),
        ("5", "Entity extraction accuracy", "≥ 85%"),
        ("6", "End-to-end query latency", "< 5 seconds"),
        ("7", "Estimated monthly cost at 500TB production", "Within budget"),
        ("8", "GovCloud deployment — all services available", "Pass/Fail"),
    ]:
        add_table_row(tc, [num, crit, target, ""])

    # ---- Budget ----
    doc.add_heading("Budget Estimate", level=2)
    tb = doc.add_table(rows=1, cols=3)
    tb.style = "Table Grid"
    tb.rows[0].cells[0].text = "Category"
    tb.rows[0].cells[1].text = "Estimated Cost"
    tb.rows[0].cells[2].text = "Notes"
    style_header_row(tb)
    for cat, cost, notes in [
        ("ProServe / Implementation", "$120,000 - $140,000", "Architecture, pipeline build, AI integration, testing"),
        ("AWS Services (pilot period)", "$40,000 - $60,000", "Neptune, OpenSearch, Aurora, Bedrock, S3, Lambda, Step Functions"),
        ("Bedrock (entity extraction)", "$15,000 - $25,000", "Claude Haiku at ~$0.003/doc for pilot subset"),
        ("Textract (OCR if needed)", "$5,000 - $15,000", "Only for scanned/image PDFs at ~$1.50/1000 pages"),
        ("Total", "~$200,000", "Combined ProServe + AWS consumption"),
    ]:
        add_table_row(tb, [cat, cost, notes])

    # ---- Risks ----
    doc.add_heading("Risks & Mitigations", level=2)
    tr = doc.add_table(rows=1, cols=3)
    tr.style = "Table Grid"
    tr.rows[0].cells[0].text = "Risk"
    tr.rows[0].cells[1].text = "Impact"
    tr.rows[0].cells[2].text = "Mitigation"
    style_header_row(tr)
    for risk, impact, mit in [
        ("High % of scanned PDFs requiring OCR", "Increased Textract cost + processing time", "Data assessment in Week 2; Textract only for non-readable; budget buffer"),
        ("500TB exceeds pilot budget", "Scope creep", "Start with representative 10-50TB subset; project production costs separately"),
        ("OpenSearch managed loader not in GovCloud", "Custom loader maintenance", "Build aligned with service team code; easy migration when managed service launches"),
        ("Neptune graph size at 500TB scale", "Performance degradation", "Partition by case; use Neptune Analytics for cross-case; benchmark during pilot"),
        ("GovCloud service availability gaps", "Architecture changes needed", "Validate all services in GovCloud during Week 1; identify alternatives early"),
        ("Data sensitivity / ATO requirements", "Delayed pilot start", "Engage security team in Week 1; use existing ATO boundary if possible"),
    ]:
        add_table_row(tr, [risk, impact, mit])

    # ---- Evaluation ----
    doc.add_heading("Evaluation Criteria", level=2)
    te = doc.add_table(rows=1, cols=4)
    te.style = "Table Grid"
    te.rows[0].cells[0].text = "Criteria"
    te.rows[0].cells[1].text = "Value"
    te.rows[0].cells[2].text = "Simplicity/Strength"
    te.rows[0].cells[3].text = "PoC Status"
    style_header_row(te)
    for crit in [
        "Ingestion Performance (docs/hour)",
        "Search Relevance (semantic + keyword)",
        "Entity Extraction Quality",
        "Cross-Case Pattern Discovery",
        "AI Briefing / Summary Quality",
        "Knowledge Graph Visualization",
        "Scalability Path to 500TB",
        "Cost Model at Production Scale",
        "GovCloud Compatibility",
        "Integration with Existing Tools",
        "User Experience / Analyst Workflow",
    ]:
        add_table_row(te, [crit, "", "", ""])

    # ---- Next Steps ----
    doc.add_heading("Next Steps", level=2)
    for step in [
        "Complete CIO Planning Call Questionnaire",
        "Finalize pilot dataset selection and data assessment",
        "Confirm budget allocation and ProServe engagement model",
        "Schedule technical kickoff with engineering team",
        "Begin GovCloud environment setup",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("© 2026, Amazon Web Services, Inc. or its affiliates. All rights reserved.")

    path = os.path.join(OUT_DIR, "Investigative-Intelligence-PoC-Success-Plan.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    q_path = make_questionnaire()
    s_path = make_success_plan()
    print(f"\nDone. Files at:\n  {q_path}\n  {s_path}")
