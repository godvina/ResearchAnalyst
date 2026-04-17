"""Convert the data governance markdown to a Word document."""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT = "docs/Investigative-Intelligence-Data-Governance.md"
OUTPUT = "docs/Investigative-Intelligence-Data-Governance.docx"

doc = Document()

# Styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# === TITLE PAGE ===
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Investigative Intelligence Platform")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
run.bold = True

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("Data Governance & Access Control")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x48, 0xBB, 0x78)

doc.add_paragraph()

for line_text in [
    "Document Version: 1.0",
    "Date: April 14, 2026",
    "Author: Emerging Tech Solutions",
    "Classification: AWS Internal / NDA",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Encryption · Access Control · Audit Trail · Compliance\n"
    "Aurora · S3 · Bedrock · VPC PrivateLink · KMS"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
run.italic = True

doc.add_page_break()

# === TABLE OF CONTENTS ===
toc_heading = doc.add_heading("Table of Contents", level=1)
for run in toc_heading.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

toc_entries = [
    "1. Data Residency & Isolation",
    "2. Encryption",
    "3. Access Control Models",
    "   Option A: Label-Based Access Control",
    "   Option B: Row-Level Security (RLS)",
    "   Option C: Hybrid (Recommended)",
    "4. Discovery Questions for the Customer",
    "   Organizational Structure",
    "   Classification & Sensitivity",
    "   Identity & Authentication",
    "   Compliance & Audit",
    "   Data Lifecycle",
    "5. Audit Trail Architecture",
    "6. AI Data Governance",
    "7. Infrastructure Security Controls",
    "8. Compliance Alignment",
    "9. Recommendation",
]

for entry_text in toc_entries:
    p = doc.add_paragraph()
    if entry_text.startswith("   "):
        run = p.add_run("     " + entry_text.strip())
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    else:
        run = p.add_run(entry_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        run.bold = True

doc.add_page_break()


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h


def add_table_from_lines(lines):
    """Parse markdown table lines into a Word table."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    header = rows[0]
    data_rows = [r for r in rows[1:] if not all(set(c) <= {"-", " ", ":"} for c in r)]

    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell_text in enumerate(header):
        table.rows[0].cells[i].text = cell_text
    for r_idx, row in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < len(table.columns):
                table.rows[r_idx + 1].cells[c_idx].text = cell_text


# === PARSE MARKDOWN BODY ===
with open(INPUT, "r", encoding="utf-8") as f:
    content = f.read()

# Skip the front-matter (title, version, date lines) — we already have a title page
lines = content.split("\n")
# Find where section 1 starts
start_idx = 0
for idx, line in enumerate(lines):
    if line.startswith("## 1."):
        start_idx = idx
        break

i = start_idx
table_buffer = []
in_table = False
in_code = False

while i < len(lines):
    line = lines[i]

    # Code blocks
    if line.strip().startswith("```"):
        if in_code:
            in_code = False
            i += 1
            continue
        else:
            in_code = True
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            in_code = False
            i += 1
            continue

    # Tables
    if "|" in line and not in_code:
        if not in_table:
            in_table = True
            table_buffer = []
        table_buffer.append(line)
        i += 1
        continue
    elif in_table:
        add_table_from_lines(table_buffer)
        doc.add_paragraph()
        table_buffer = []
        in_table = False

    # Headings
    if line.startswith("## "):
        add_heading(line[3:].strip(), level=1)
    elif line.startswith("### "):
        add_heading(line[4:].strip(), level=2)
    elif line.startswith("#### "):
        add_heading(line[5:].strip(), level=3)
    elif line.startswith("---"):
        doc.add_paragraph()
    elif line.strip() == "":
        pass
    elif re.match(r"^\s*- ", line):
        # Bullets — including indented sub-bullets
        bullet_text = re.sub(r"^\s*- ", "", line).strip()
        bullet_text = re.sub(r"\*\*(.+?)\*\*", r"\1", bullet_text)
        bullet_text = re.sub(r"`(.+?)`", r"\1", bullet_text)
        indent = len(line) - len(line.lstrip())
        if indent >= 2:
            p = doc.add_paragraph(bullet_text, style="List Bullet 2")
        else:
            p = doc.add_paragraph(bullet_text, style="List Bullet")
    elif re.match(r"^\d+\.\s", line):
        # Numbered list — strip number, handle bold question text
        text = re.sub(r"^\d+\.\s*", "", line).strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        p = doc.add_paragraph(text, style="List Number")
    elif line.startswith("**") and line.endswith("**"):
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*").strip())
        run.bold = True
    else:
        clean = line.strip()
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)
        clean = re.sub(r"`(.+?)`", r"\1", clean)
        if clean:
            doc.add_paragraph(clean)

    i += 1

# Flush remaining table
if in_table and table_buffer:
    add_table_from_lines(table_buffer)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
