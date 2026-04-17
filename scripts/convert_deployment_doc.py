"""Convert the deployment architecture markdown to a Word document."""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT = "docs/Investigative-Intelligence-Deployment-Architecture.md"
OUTPUT = "docs/Investigative-Intelligence-Deployment-Architecture.docx"

doc = Document()

# Styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# === TITLE PAGE ===
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Investigative Intelligence Platform")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
run.bold = True

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("Multi-Environment Deployment Architecture")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x48, 0xBB, 0x78)

doc.add_paragraph()

for line_text in [
    "Document Version: 1.0",
    "Date: April 12, 2026",
    "Author: Emerging Tech Solutions",
    "Classification: AWS Internal / NDA",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Serverless AI-Powered Investigative Analysis\nAurora · Neptune · Bedrock · OpenSearch · Lambda · Step Functions")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
run.italic = True

doc.add_page_break()

# === TABLE OF CONTENTS ===
toc_heading = doc.add_heading("Table of Contents", level=1)
for run in toc_heading.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

# Insert a manual TOC with section titles and page references
toc_entries = [
    ("1. Platform Architecture Overview", 1),
    ("   1.1 AWS Services Used", 1),
    ("   1.2 Architecture Diagram", 1),
    ("   1.3 Ingestion Pipeline (Step Functions)", 1),
    ("2. Deployment Tiers", 2),
    ("   2.1 Tier 1: Demo (Isengard Standard)", 2),
    ("   2.2 Tier 2: GovCloud Test (Isengard GovCloud)", 2),
    ("   2.3 Tier 3: Customer Production (DOJ GovCloud)", 2),
    ("3. Config-Driven Deployment System", 2),
    ("   3.1 How It Works", 2),
    ("   3.2 What the Config Controls", 2),
    ("   3.3 Graceful Degradation", 2),
    ("4. GovCloud Considerations", 3),
    ("   4.1 GovCloud vs Commercial — Key Differences", 3),
    ("   4.2 Federal Security Controls", 3),
    ("   4.3 FedRAMP Bedrock Model Availability", 3),
    ("   4.4 DOJ-Specific Considerations", 3),
    ("5. Deployment Runbook", 3),
    ("   5.1 Prerequisites", 3),
    ("   5.2 Tier 1: Demo Deployment (10 Steps)", 3),
    ("   5.3 Tier 2: GovCloud Deployment", 3),
    ("   5.4 CloudFormation Console Deployment", 3),
    ("6. Modular CDK Architecture", 4),
    ("7. Scaling for 500TB Ingestion", 4),
    ("8. Next Steps", 4),
    ("9. Code Compliance & Federal Deployment Readiness", 5),
    ("   9.1 Re-Deployment After Changes", 5),
    ("   9.2 Software Supply Chain Assessment", 5),
    ("   9.3 Federal Code Approval Process", 5),
    ("   9.4 Authority to Operate (ATO) Considerations", 5),
    ("   9.5 Static Application Security Testing (SAST)", 5),
    ("   9.6 AI-Generated Code Disclosure", 5),
    ("   9.7 Pre-Deployment Checklist for SA Handoff", 6),
    ("   9.8 Questions to Ask the SA Before GovCloud Deployment", 6),
    ("10. Risk Register", 6),
    ("Appendix A: Deployment Config Schema", 6),
    ("Appendix B: Current Platform Capabilities", 6),
]

for entry_text, _ in toc_entries:
    p = doc.add_paragraph()
    if entry_text.startswith("   "):
        # Sub-entry — indented
        run = p.add_run("     " + entry_text.strip())
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    else:
        run = p.add_run(entry_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        run.bold = True

doc.add_page_break()

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h

def add_table_from_lines(lines):
    """Parse markdown table lines into a Word table."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    # Skip separator row (row[1] with dashes)
    header = rows[0]
    data_rows = [r for r in rows[1:] if not all(set(c) <= {"-", " ", ":"} for c in r)]
    
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    # Header
    for i, cell_text in enumerate(header):
        table.rows[0].cells[i].text = cell_text
    # Data
    for r_idx, row in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < len(table.columns):
                table.rows[r_idx + 1].cells[c_idx].text = cell_text

with open(INPUT, "r", encoding="utf-8") as f:
    content = f.read()

lines = content.split("\n")
i = 0
table_buffer = []
in_table = False
in_code = False

while i < len(lines):
    line = lines[i]
    
    # Code blocks
    if line.strip().startswith("```"):
        if in_code:
            in_code = False
            i += 1
            continue
        else:
            in_code = True
            lang = line.strip().replace("```", "").strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            in_code = False
            i += 1
            continue
    
    # Tables
    if "|" in line and not in_code:
        if not in_table:
            in_table = True
            table_buffer = []
        table_buffer.append(line)
        i += 1
        continue
    elif in_table:
        add_table_from_lines(table_buffer)
        doc.add_paragraph()  # spacing
        table_buffer = []
        in_table = False
    
    # Headings
    if line.startswith("# "):
        add_heading(line[2:].strip(), level=1)
    elif line.startswith("## "):
        add_heading(line[3:].strip(), level=2)
    elif line.startswith("### "):
        add_heading(line[4:].strip(), level=3)
    elif line.startswith("---"):
        doc.add_paragraph()  # page break substitute
    elif line.strip() == "":
        pass  # skip blank lines
    elif line.startswith("- "):
        p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
    elif re.match(r"^\d+\.", line):
        text = re.sub(r"^\d+\.\s*", "", line).strip()
        p = doc.add_paragraph(text, style="List Number")
    elif line.startswith("**") and line.endswith("**"):
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*").strip())
        run.bold = True
    else:
        # Regular paragraph — strip markdown bold/italic
        clean = line.strip()
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)
        if clean:
            doc.add_paragraph(clean)
    
    i += 1

# Flush remaining table
if in_table and table_buffer:
    add_table_from_lines(table_buffer)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
